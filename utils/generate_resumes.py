# utils/generate_resumes.py - Programmatic generation of 10 distinct, detailed ATS-friendly resumes.

import os
from docx import Document
from utils.logger import logger

def add_header(doc, name, email, phone, location, title):
    doc.add_heading(name, 0)
    p = doc.add_paragraph()
    p.add_run(f"{title} | {location} | {phone} | {email}").italic = True

def add_section(doc, heading_text, content_list, is_bullet=True):
    doc.add_heading(heading_text, level=1)
    for item in content_list:
        if is_bullet:
            doc.add_paragraph(item, style='List Bullet')
        else:
            doc.add_paragraph(item)

def generate_trainee_tool_die(path):
    doc = Document()
    add_header(doc, "ROHAN SHARMA", "rohan.sharma@email.com", "+91-8899001122", "Gurgaon, India", "Trainee Tool & Die Maker")
    
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(
        "Ambitious and disciplined entry-level technical professional with a solid foundation in mechanical engineering concepts, "
        "workshop calculations, and precision tooling. Eager to assist senior tool makers on the assembly shop floor to manufacture, "
        "spot-align, and maintain high-precision jigs, fixtures, and metal-stamping dies. Fast learner with strict compliance to tool room safety rules."
    )
    
    doc.add_heading("Education & Credentials", level=1)
    doc.add_paragraph("Diploma in Tool & Die Making (3-Year Course)\nIndustrial Training Institute (ITI), Delhi | Graduated 2025 (84% Score)")
    
    doc.add_heading("Core Academic Projects", level=1)
    p = doc.add_paragraph()
    p.add_run("Fabrication of a Manual Press Tool Assembly:\n").bold = True
    p.add_run(
        "Designed and fabricated a simple blanking press tool as a group final project. "
        "Executed turning, milling, and surface grinding operations on raw steel blocks. "
        "Calculated shear clearances and aligned punch-and-die sets to clear scrap blanking defects."
    )

    add_section(doc, "Technical Competencies", [
        "Machining: Basic operation of conventional lathes, drill presses, and shaping machines.",
        "Metrology: Precise handling of workshop tools including micrometers, verniers, and dial test indicators.",
        "Engineering Graphics: Interpretation of orthographic engineering blueprints and CAD layouts.",
        "Safety Regulations: Strict adherence to personal protective equipment (PPE) and machine emergency shutdown procedures."
    ], is_bullet=True)
    
    doc.save(path)

def generate_apprentice_tool_maker(path):
    doc = Document()
    add_header(doc, "VIKRAM JOSHI", "vikram.joshi@email.com", "+91-7788990011", "Ahmedabad, India", "Apprentice Tool Maker")
    
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(
        "Motivated apprentice machinist seeking structured practical exposure in mold assembly and tool room setups. "
        "Possesses a strong mechanical aptitude, willingness to follow precise work orders under supervision, "
        "and solid hands-on experience in manual milling and precision thread tapping."
    )
    
    doc.add_heading("Education", level=1)
    doc.add_paragraph("ITI Machinist Certification\nGujarat Vocational Training Board | Completed 2025")
    
    doc.add_heading("Practical Training Logbook", level=1)
    p = doc.add_paragraph()
    p.add_run("Apprentice Workshop Practices (2024 - 2025):\n").bold = True
    p.add_run(
        "Performed daily cleaning, lubrication, and basic setup adjustments on surface grinders and drill machines. "
        "Trained in manual die assembly, screw thread tapping, filing, and finishing core blocks. "
        "Inspected machined test pieces under the guidance of supervisor tool room technicians."
    )

    add_section(doc, "Key Skills learned", [
        "Handling vernier calipers, bevel protractors, and height gauges.",
        "Identifying various grades of tool steels (D2, H13, O1).",
        "Understanding slide clearance allowances in guide pillars and bushings.",
        "Applying safety protocols during crane movements of heavy tooling cores."
    ], is_bullet=True)
    
    doc.save(path)

def generate_die_maintenance_trainee(path):
    doc = Document()
    add_header(doc, "SANDEEP KUMAR", "sandeep.k@email.com", "+91-9900112233", "Pune, India", "Die Maintenance Trainee")
    
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(
        "Entry-level mechanical technician specializing in the maintenance and repair of sheet-metal stamping dies. "
        "Passionate about troubleshooting die-wear issues, performing assembly clearances, and cleaning mold elements. "
        "Familiar with press operations and structural maintenance schedules."
    )
    
    doc.add_heading("Education", level=1)
    doc.add_paragraph("ITI Fitter / Tool & Die Trainee\nState Board of Technical Education, Maharashtra | Completed 2024")
    
    add_section(doc, "Practical Skills & Exposure", [
        "Die inspection: Checking punches and dies for chipping, wear, or misalignment.",
        "Lubrication & Assembly: Assembling pillars, spring cylinders, and applying lubricants.",
        "Troubleshooting: Assisting in replacing worn ejector pins, guide bushings, and damaged springs.",
        "Record keeping: Maintaining clear documentation logs of maintenance cycles and tool room metrics."
    ], is_bullet=True)
    
    doc.save(path)

def generate_mold_shop_helper(path):
    doc = Document()
    add_header(doc, "MANOJ YADAV", "manoj.yadav@email.com", "+91-8811223344", "Noida, India", "Mold Shop Helper")
    
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(
        "Dutiful and physically active shop assistant eager to support mold makers, CNC machinists, and engineers. "
        "Experienced in safe material handling, cleaning core/cavity parts, polishing metal surfaces, and organizing inventory."
    )
    
    doc.add_heading("Education & Training", level=1)
    doc.add_paragraph("10th Standard / Vocational Welding Certificate\nUP Secondary Board | Completed 2023")
    
    add_section(doc, "Responsibilities & Capabilities", [
        "Cleaning injection mold surfaces using specialized solvent cleaners and rust preventatives.",
        "Polishing mold cavities manually using compound pastes to achieve optical finish.",
        "Transporting heavy raw steel blocks and components safely across CNC machines using mechanical cranes.",
        "Maintaining tool cabinets, checklist boards, and scrap metal disposal bins."
    ], is_bullet=True)
    
    doc.save(path)

def generate_progressive_die_maker(path):
    doc = Document()
    add_header(doc, "HARINDER SINGH", "harinder.s@email.com", "+91-9555667788", "Faridabad, India", "Progressive Die Maker")
    
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(
        "Senior Tool & Die Maker with over 6 years of expertise specializing in complex, multi-stage progressive stamping dies. "
        "Proven track record of building, timing, and troubleshooting progressive layouts to ensure high-volume blanking, "
        "bending, and forming repeatability within 0.01mm tolerances. Deep understanding of sheet-metal behavior under high tonnage presses."
    )
    
    doc.add_heading("Professional Experience", level=1)
    p = doc.add_paragraph()
    p.add_run("Senior Die Maker - Escorts Auto Component Division (2020 - Present)\n").bold = True
    p.add_run(
        "Manufactured and timed 5-stage progressive stamping dies for automotive panel components. "
        "Interpreted 2D/3D die assembly designs, strip layouts, and clearances. "
        "Spotted and aligned punches and dies on 200-ton hydraulic press machines. "
        "Reduced initial part trial defects by 25% through precision adjustments of guide lifters and pilot positions."
    )
    
    doc.add_heading("Education", level=1)
    doc.add_paragraph("Diploma in Tool & Die Making (4-Year Course)\nNTTF (Nettur Technical Training Foundation) | Graduated 2018")
    
    add_section(doc, "Key Skills", [
        "Stamping tooling: Multi-stage progressive dies, drawing dies, compound blanking dies.",
        "Machining: VMC G-code path adjustments, precision surface grinding, cylindrical turning.",
        "Metrology: Optical comparators, coordinate measuring machines (CMM), micrometers.",
        "Materials: D2, H13, Vanadis tool steels, heat treatment guidelines."
    ], is_bullet=True)
    
    doc.save(path)

def generate_injection_mold_maker(path):
    doc = Document()
    add_header(doc, "ANIL DESHMUKH", "anil.d@email.com", "+91-9444556677", "Aurangabad, India", "Injection Mold Maker")
    
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(
        "Skilled Mold Maker with 5 years of experience in the design, assembly, and testing of plastic injection molds. "
        "Expert in finishing core and cavity plates, setting up water cooling paths, and aligning slide cores. "
        "Adept at manual high-precision grinding and fitting operations."
    )
    
    doc.add_heading("Professional Experience", level=1)
    p = doc.add_paragraph()
    p.add_run("Mold Maker - Supreme Plastics Ltd (2021 - Present)\n").bold = True
    p.add_run(
        "Assembled multi-cavity injection molds, hot runner blocks, and manifold assemblies. "
        "Managed polishing of mold cores to achieve high mirror finish. "
        "Drilled and routed cooling channels, checked seal gaskets, and tested slide-core limits. "
        "Collaborated with molding production technicians during trials to adjust draft offsets and vent heights."
    )
    
    doc.add_heading("Education", level=1)
    doc.add_paragraph("Diploma in Mold Technology\nCIPET (Central Institute of Petrochemicals Engineering & Technology) | Completed 2020")
    
    add_section(doc, "Skills & Tools", [
        "Assembly of hot-runner manifolds and ejector plate components.",
        "Grinding and polishing of mirror-finished cavity surfaces.",
        "Familiar with hydraulic core pull systems and slide positioning logic.",
        "Usage of height masters, bore gauges, and slip gauge block configurations."
    ], is_bullet=True)
    
    doc.save(path)

def generate_cnc_machinist(path):
    doc = Document()
    add_header(doc, "KARTIK MEHTA", "kartik.mehta@email.com", "+91-9111223344", "Bangalore, India", "CNC Machinist")
    
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(
        "VMC & CNC Programmer/Machinist with 4 years of experience writing G-code, setting work offsets, "
        "and cutting precision tooling components. Expert in programming Fanuc and Siemens controls, "
        "maximizing tool life, and interpreting mechanical GD&T engineering prints."
    )
    
    doc.add_heading("Professional Experience", level=1)
    p = doc.add_paragraph()
    p.add_run("CNC Operator & Programmer - Bharat Forge Ltd (2022 - Present)\n").bold = True
    p.add_run(
        "Programmed 3-axis Vertical Machining Centers (VMC) for core mold cavity blocks. "
        "Determined feed rates, cutting speeds, and selected appropriate carbide cutters for D2 steel. "
        "Performed work offset settings (G54-G59) and adjusted tool wear offsets. "
        "Inspected final dimensions using digital vernier calipers and height gauges to ensure 0.01mm compliance."
    )
    
    doc.add_heading("Education", level=1)
    doc.add_paragraph("Diploma in Mechanical Engineering\nState Board of Technical Education, Karnataka | Completed 2021")
    
    add_section(doc, "Technical Skills", [
        "Programming: Manual G-Code/M-Code, basic Mastercam layouts.",
        "Controllers: Fanuc, Siemens Sinumerik, Haas controllers.",
        "Operations: Face milling, profile contouring, pocketing, deep-hole drilling.",
        "Inspections: Dial bore indicators, thread gauges, digital micrometers."
    ], is_bullet=True)
    
    doc.save(path)

def generate_edm_operator(path):
    doc = Document()
    add_header(doc, "BALAJI NAIDU", "balaji.n@email.com", "+91-9000887766", "Coimbatore, India", "EDM Operator")
    
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(
        "Specialized EDM Operator with 3 years of hands-on experience in Wire Cut EDM and Sinker EDM machining. "
        "Adept at preparing copper/graphite electrodes, programming wire profile paths, and cutting complex cavity shapes "
        "in hardened die blocks within 5-micron tolerances."
    )
    
    doc.add_heading("Professional Experience", level=1)
    p = doc.add_paragraph()
    p.add_run("EDM Machinist - Kovai Tool Room Engineering (2023 - Present)\n").bold = True
    p.add_run(
        "Operated Mitsubishi Wire Cut EDM and Sinker EDM systems. "
        "Fitted workpieces and dialed coordinates to prepare spark erosion setups. "
        "Programmed complex geometries to cut progressive die slots and gear templates. "
        "Inspected surface finish index and verified dimensions against GD&T specification prints."
    )
    
    doc.add_heading("Education", level=1)
    doc.add_paragraph("ITI Machinist Certification\nTamil Nadu Board of Technical Training | Completed 2022")
    
    add_section(doc, "Technical Competencies", [
        "Wire EDM programming: Wire threading, alignment, feed calculations.",
        "Sinker EDM practices: Electrode alignment, flush pressure settings, amperage control.",
        "Metrology checks: Profile projectors, micrometer calipers, toolmaker microscopes.",
        "Materials: Hardened tool steel (D2, H13 at HRC 58-62), tungsten carbides."
    ], is_bullet=True)
    
    doc.save(path)

def generate_die_design_engineer(path):
    doc = Document()
    add_header(doc, "ARUN PRASAD", "arun.prasad@email.com", "+91-9333445566", "Chennai, India", "Die Design Engineer")
    
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(
        "Mechanical Design Engineer with 4 years of experience specializing in sheet-metal press tools, progressive dies, "
        "and structural assembly fixtures. Fluent in creating strip layouts, calculation parameters (cutting force, strip utilization), "
        "and utilizing SolidWorks and AutoCAD for ASME Y14.5 GD&T compliance."
    )
    
    doc.add_heading("Professional Experience", level=1)
    p = doc.add_paragraph()
    p.add_run("Design Engineer - Lucas TVS Ltd (2022 - Present)\n").bold = True
    p.add_run(
        "Designed progressive, compound, and draw dies using SolidWorks. "
        "Optimized strip layouts to reduce metal scrap waste by 12%. "
        "Calculated tonnage limits, spring pressures, and punch load distributions. "
        "Assisted tool room technicians during initial component press trials to correct clearance splits."
    )
    
    doc.add_heading("Education", level=1)
    doc.add_paragraph("B.E. in Mechanical Engineering\nAnna University, Chennai | Completed 2021 (CGPA: 8.2)")
    
    add_section(doc, "Technical Expertise", [
        "Design Software: SolidWorks (Sheetmetal), AutoCAD, CATIA V5.",
        "Engineering Calculations: Press tonnage, strip layouts, blank size calculations, spring selection.",
        "Standards: GD&T, tolerance stack-up analysis, ISO standards.",
        "Fixture Design: Checking fixtures, welding jigs, machining clamps."
    ], is_bullet=True)
    
    doc.save(path)

def generate_mold_design_engineer(path):
    doc = Document()
    add_header(doc, "NEHA KULKARNI", "neha.k@email.com", "+91-9222334455", "Bangalore, India", "Mold Design Engineer")
    
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(
        "Analytical Mold Design Engineer with 5 years of experience creating core-and-cavity layouts for plastic injection "
        "moulds. Skilled in running MoldFlow thermal simulations, designing hot runner systems, and selecting core gating. "
        "Strong understanding of draft clearances and material shrinkage coefficients."
    )
    
    doc.add_heading("Professional Experience", level=1)
    p = doc.add_paragraph()
    p.add_run("Design Engineer - Tata Autocomp Systems (2021 - Present)\n").bold = True
    p.add_run(
        "Created 3D mold CAD models for automotive interior trims using Siemens NX. "
        "Designed gate positions, cooling channel paths, and mechanical ejector plate systems. "
        "Conducted MoldFlow shrink, warp, and temperature simulations to resolve cosmetic sink marks. "
        "Collaborated with CNC tooling machinists to verify CAD model dimensions."
    )
    
    doc.add_heading("Education", level=1)
    doc.add_paragraph("B.Tech in Plastics Technology\nCIPET, Ahmedabad | Completed 2020 (Gold Medalist)")
    
    add_section(doc, "Core Skills", [
        "CAD systems: Siemens NX, SolidWorks, AutoCAD.",
        "Mold Flow Analysis: Fill analysis, pack checks, cooling calculations.",
        "Tooling concepts: Cavity splits, sliders, angle pin actions, lifter designs.",
        "Material shrinkages: ABS, Polypropylene, Polycarbonate, Nylon parameters."
    ], is_bullet=True)
    
    doc.save(path)

def main():
    logger.info("Initializing 10 Resumes Generation...")
    target_dir = os.path.join("data", "resumes")
    os.makedirs(target_dir, exist_ok=True)
    
    resumes = {
        "resume_trainee_tool_die.docx": generate_trainee_tool_die,
        "resume_apprentice_tool_maker.docx": generate_apprentice_tool_maker,
        "resume_die_maintenance_trainee.docx": generate_die_maintenance_trainee,
        "resume_mold_shop_helper.docx": generate_mold_shop_helper,
        "resume_progressive_die_maker.docx": generate_progressive_die_maker,
        "resume_injection_mold_maker.docx": generate_injection_mold_maker,
        "resume_cnc_machinist.docx": generate_cnc_machinist,
        "resume_edm_operator.docx": generate_edm_operator,
        "resume_die_design_engineer.docx": generate_die_design_engineer,
        "resume_mold_design_engineer.docx": generate_mold_design_engineer
    }
    
    for filename, generator in resumes.items():
        file_path = os.path.join(target_dir, filename)
        if not os.path.exists(file_path):
            generator(file_path)
            
    print("Successfully generated 10 specialized tool-room resumes in data/resumes/")

if __name__ == "__main__":
    main()
